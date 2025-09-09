"""
Web API接口层 - 为现有的文献分析系统提供HTTP接口
不修改任何现有代码，只是添加API包装
"""
from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS
import os
import json
import uuid
import threading
from datetime import datetime
from werkzeug.utils import secure_filename
import time
from io import BytesIO

# 导入现有的核心模块
from core.citation_checker import CitationChecker
from config.config_manager import ConfigManager
from utils.document_parser import DocumentParser  # 新增的文档解析器
from utils.position_tracker import PositionTracker  # 新增的位置跟踪器

# 导入PDF导出所需模块
from weasyprint import HTML, CSS

# 导入Word导出所需模块
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
CORS(app)  # 允许跨域请求

# 配置
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# 确保上传目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 全局存储 - 存储分析任务
analysis_tasks = {}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

class AnalysisTask:
    """分析任务类"""
    def __init__(self, task_id, doc_path, analysis_mode):
        self.task_id = task_id
        self.doc_path = doc_path
        self.analysis_mode = analysis_mode
        self.status = 'pending'  # pending, running, completed, error
        self.progress = {'processed': 0, 'total': 0, 'percentage': 0}
        self.result = None
        self.error = None
        self.created_at = datetime.now()
        self.checker = None
        self.document_parser = None
        self.position_tracker = None

@app.route('/api/upload', methods=['POST'])
def upload_document():
    """上传文档接口"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': '没有文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '没有选择文件'}), 400
        
        if file and allowed_file(file.filename):
            # 生成唯一的任务ID
            task_id = str(uuid.uuid4())
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{task_id}_{filename}")
            file.save(file_path)
            
            # 创建分析任务
            task = AnalysisTask(task_id, file_path, 'subjective')  # 默认主观模式
            analysis_tasks[task_id] = task
            
            # 初始化文档解析器
            task.document_parser = DocumentParser(file_path)
            task.position_tracker = PositionTracker(file_path)
            
            return jsonify({
                'task_id': task_id,
                'filename': filename,
                'message': '文档上传成功'
            })
        else:
            return jsonify({'error': '不支持的文件格式，仅支持.docx文件'}), 400
            
    except Exception as e:
        return jsonify({'error': f'上传失败: {str(e)}'}), 500

@app.route('/api/document/<task_id>/preview', methods=['GET'])
def get_document_preview(task_id):
    """获取文档预览"""
    try:
        if task_id not in analysis_tasks:
            return jsonify({'error': '任务不存在'}), 404
        
        task = analysis_tasks[task_id]
        
        # 使用文档解析器获取格式化的HTML内容
        html_content = task.document_parser.to_html()
        paragraphs = task.document_parser.get_paragraphs_with_positions()
        
        return jsonify({
            'task_id': task_id,
            'html_content': html_content,
            'paragraphs': paragraphs,
            'total_paragraphs': len(paragraphs)
        })
        
    except Exception as e:
        return jsonify({'error': f'获取预览失败: {str(e)}'}), 500

@app.route('/api/document/<task_id>/analyze', methods=['POST'])
def start_analysis(task_id):
    """开始分析"""
    try:
        if task_id not in analysis_tasks:
            return jsonify({'error': '任务不存在'}), 404
        
        task = analysis_tasks[task_id]
        
        if task.status == 'running':
            return jsonify({'error': '分析已在进行中'}), 400
        
        # 获取分析模式
        data = request.get_json()
        if data and 'analysis_mode' in data:
            task.analysis_mode = data['analysis_mode']
        
        # 启动分析线程
        analysis_thread = threading.Thread(
            target=run_analysis, 
            args=(task,)
        )
        analysis_thread.start()
        
        return jsonify({
            'task_id': task_id,
            'message': '分析已开始',
            'analysis_mode': task.analysis_mode
        })
        
    except Exception as e:
        return jsonify({'error': f'启动分析失败: {str(e)}'}), 500

def run_analysis(task):
    """运行分析的后台任务"""
    try:
        task.status = 'running'
        print(f"开始分析任务: {task.task_id}")
        
        # 根据分析模式选择配置文件
        if task.analysis_mode == "subjective":
            config_path = os.path.join(os.path.dirname(__file__), "config", "config_subjective.json")
        elif task.analysis_mode == "quick":
            config_path = os.path.join(os.path.dirname(__file__), "config", "config_quick.json")
        else:
            config_path = os.path.join(os.path.dirname(__file__), "config", "config_full.json")
        
        # 使用现有的CitationChecker进行分析
        task.checker = CitationChecker(task.doc_path, config_path)
        
        # 重写进度更新方法以更新任务状态
        original_update_progress = task.checker._update_progress
        def update_task_progress():
            task.progress = {
                'processed': task.checker.processed_citations,
                'total': task.checker.total_citations,
                'percentage': int((task.checker.processed_citations / task.checker.total_citations) * 100) if task.checker.total_citations > 0 else 0
            }
            print(f"进度更新: {task.progress['processed']}/{task.progress['total']} ({task.progress['percentage']}%)")
            original_update_progress()
        task.checker._update_progress = update_task_progress
        
        # 生成报告
        print("开始生成报告...")
        report = task.checker.generate_report()
        print("报告生成完成")
        
        # 解析报告并提取问题信息
        print("开始提取问题信息...")
        problems = extract_problems_from_report(task.checker, task.position_tracker)
        print(f"问题提取完成，发现 {len(problems)} 个问题")
        
        # 确保进度为100%
        task.progress = {
            'processed': task.checker.total_citations,
            'total': task.checker.total_citations,
            'percentage': 100
        }
        
        task.result = {
            'report': report,
            'problems': problems,
            'analysis_mode': task.analysis_mode,
            'completed_at': datetime.now().isoformat()
        }
        task.status = 'completed'
        print(f"分析任务完成: {task.task_id}")
        
    except Exception as e:
        task.status = 'error'
        task.error = str(e)
        print(f"分析任务失败: {task.task_id}, 错误: {e}")
        import traceback
        traceback.print_exc()

def extract_problems_from_report(checker, position_tracker):
    """从分析结果中提取问题信息"""
    problems = []
    problem_id = 1
    
    # 1. 缺失引用问题
    for citation in checker.missing_citations:
        positions = position_tracker.find_citation_positions(citation)
        for pos in positions:
            problems.append({
                'id': problem_id,
                'type': 'missing_citation',
                'severity': 'error',
                'citation': citation,
                'description': f'引用 {citation} 未在参考文献中找到',
                'position': pos,
                'color': '#ff4d4f'  # 红色
            })
            problem_id += 1
    
    # 2. 未使用的参考文献
    unused_refs = checker.check_unused_references()
    for ref in unused_refs:
        ref_positions = position_tracker.find_reference_positions(ref['text'])
        for pos in ref_positions:
            problems.append({
                'id': problem_id,
                'type': 'unused_reference',
                'severity': 'warning',
                'reference': ref['text'],
                'description': f'参考文献未被正文引用',
                'position': pos,
                'color': '#faad14'  # 黄色
            })
            problem_id += 1
    
    # 3. 不相关引用问题 (需要解析AI分析结果)
    for citation in checker.citations:
        if citation not in checker.missing_citations:
            # 这里需要解析AI分析结果，判断是否不相关
            # 由于现有代码的分析结果在generate_report中，我们需要重新分析
            context = checker.find_context_around_citation(citation)
            if checker.analysis_mode == "quick":
                analysis = checker.analyze_citation_relevance_quick(citation, context)
            elif checker.analysis_mode == "subjective":
                analysis = checker.analyze_citation_relevance_subjective(citation, context)
            else:
                analysis = checker.analyze_citation_relevance(citation, context)
            
            # 简单的文本分析判断是否不相关
            if '不相关' in analysis or 'not relevant' in analysis.lower():
                positions = position_tracker.find_citation_positions(citation)
                for pos in positions:
                    problems.append({
                        'id': problem_id,
                        'type': 'irrelevant_citation',
                        'severity': 'warning',
                        'citation': citation,
                        'description': f'引用 {citation} 与上下文不相关',
                        'analysis': analysis,
                        'context': context,
                        'position': pos,
                        'color': '#fa8c16'  # 橙色
                    })
                    problem_id += 1
    
    # 按文档位置排序
    problems.sort(key=lambda x: (x['position']['paragraph_index'], x['position']['start_pos']))
    
    return problems

@app.route('/api/document/<task_id>/progress', methods=['GET'])
def get_analysis_progress(task_id):
    """获取分析进度"""
    try:
        if task_id not in analysis_tasks:
            return jsonify({'error': '任务不存在'}), 404
        
        task = analysis_tasks[task_id]
        
        return jsonify({
            'task_id': task_id,
            'status': task.status,
            'progress': task.progress,
            'error': task.error
        })
        
    except Exception as e:
        return jsonify({'error': f'获取进度失败: {str(e)}'}), 500

@app.route('/api/document/<task_id>/problems', methods=['GET'])
def get_problems(task_id):
    """获取问题列表"""
    try:
        if task_id not in analysis_tasks:
            return jsonify({'error': '任务不存在'}), 404
        
        task = analysis_tasks[task_id]
        
        if task.status != 'completed':
            return jsonify({'error': '分析尚未完成'}), 400
        
        return jsonify({
            'task_id': task_id,
            'problems': task.result['problems'],
            'total_problems': len(task.result['problems']),
            'analysis_mode': task.result['analysis_mode']
        })
        
    except Exception as e:
        return jsonify({'error': f'获取问题失败: {str(e)}'}), 500

@app.route('/api/document/<task_id>/report', methods=['GET'])
def get_full_report(task_id):
    """获取完整报告"""
    try:
        if task_id not in analysis_tasks:
            return jsonify({'error': '任务不存在'}), 404
        
        task = analysis_tasks[task_id]
        
        if task.status != 'completed':
            return jsonify({'error': '分析尚未完成'}), 400
        
        return jsonify({
            'task_id': task_id,
            'report': task.result['report'],
            'analysis_mode': task.result['analysis_mode'],
            'completed_at': task.result['completed_at']
        })
        
    except Exception as e:
        return jsonify({'error': f'获取报告失败: {str(e)}'}), 500


@app.route('/api/document/<task_id>/export/pdf', methods=['GET'])
def export_report_pdf(task_id):
    """导出PDF格式报告"""
    try:
        if task_id not in analysis_tasks:
            return jsonify({'error': '任务不存在'}), 404
        
        task = analysis_tasks[task_id]
        
        if task.status != 'completed':
            return jsonify({'error': '分析尚未完成'}), 400
        
        # 获取报告HTML内容
        report_html = task.result['report']
        
        # 添加基本的HTML结构和样式
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>文献引用合规性检查报告</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }}
                h2 {{ color: #666; margin-top: 30px; }}
                h3 {{ color: #888; margin-top: 20px; }}
                ul {{ margin: 10px 0; }}
                li {{ margin: 5px 0; }}
                .context {{ background-color: #f5f5f5; padding: 10px; margin: 10px 0; border-radius: 5px; }}
                .analysis {{ background-color: #e8f4f8; padding: 10px; margin: 10px 0; border-radius: 5px; }}
            </style>
        </head>
        <body>
            <h1>文献引用合规性检查报告</h1>
            <p><strong>分析模式:</strong> {task.result['analysis_mode']}</p>
            <p><strong>完成时间:</strong> {task.result['completed_at']}</p>
            {report_html}
        </body>
        </html>
        """
        
        # 转换为PDF
        html = HTML(string=full_html)
        pdf_content = html.write_pdf()
        
        # 创建响应
        response = Response(
            pdf_content,
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename=report_{task_id}.pdf'
            }
        )
        
        return response
        
    except Exception as e:
        return jsonify({'error': f'导出PDF失败: {str(e)}'}), 500


@app.route('/api/document/<task_id>/export/word', methods=['GET'])
def export_report_word(task_id):
    """导出Word格式报告"""
    try:
        if task_id not in analysis_tasks:
            return jsonify({'error': '任务不存在'}), 404
        
        task = analysis_tasks[task_id]
        
        if task.status != 'completed':
            return jsonify({'error': '分析尚未完成'}), 400
        
        # 创建Word文档
        document = Document()
        
        # 添加标题
        document.add_heading('文献引用合规性检查报告', 0)
        
        # 添加基本信息
        document.add_paragraph(f'分析模式: {task.result["analysis_mode"]}')
        document.add_paragraph(f'完成时间: {task.result["completed_at"]}')
        
        # 添加报告内容
        report_html = task.result['report']
        
        # 简单解析HTML并转换为Word文档
        # 这里我们使用简化的方法处理HTML内容
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(report_html, 'html.parser')
        
        # 处理各个元素
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'li']):
            if element.name == 'h1':
                document.add_heading(element.get_text(), 1)
            elif element.name == 'h2':
                document.add_heading(element.get_text(), 2)
            elif element.name == 'h3':
                document.add_heading(element.get_text(), 3)
            elif element.name == 'p':
                document.add_paragraph(element.get_text())
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    document.add_paragraph(li.get_text(), style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    document.add_paragraph(li.get_text(), style='List Number')
        
        # 保存到内存
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        # 创建响应
        response = Response(
            buffer.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename=report_{task_id}.docx'
            }
        )
        
        return response
        
    except Exception as e:
        return jsonify({'error': f'导出Word失败: {str(e)}'}), 500


@app.route('/api/document/<task_id>/export', methods=['GET'])
def export_report(task_id):
    """导出报告文件（支持HTML/TXT/PDF格式）"""
    try:
        print(f"导出报告请求，任务ID: {task_id}")
        print(f"当前存储的任务: {list(analysis_tasks.keys())}")
        
        # 获取格式参数，默认为html
        format_type = request.args.get('format', 'html').lower()
        if format_type not in ['html', 'txt', 'pdf']:
            format_type = 'html'
        
        if task_id not in analysis_tasks:
            return jsonify({'error': '任务不存在'}), 404
        
        task = analysis_tasks[task_id]
        
        if task.status != 'completed':
            return jsonify({'error': '分析尚未完成'}), 400
        
        # 生成导出文件名
        from datetime import datetime
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # 根据格式类型处理
        if format_type == 'txt':
            return export_report_as_txt(task, timestamp)
        elif format_type == 'pdf':
            return export_report_as_pdf(task, timestamp)
        else:  # html
            return export_report_as_html(task, timestamp)
        
    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500


def export_report_as_html(task, timestamp):
    """导出为HTML格式"""
    filename = f"citation_report_{task.result['analysis_mode']}_{timestamp}.html"
    
    # 创建完整的HTML文件
    html_content = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>文献引用合规性检查报告</title>
    <style>
        body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        .report-header {{ text-align: center; margin-bottom: 30px; }}
        .report-meta {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
        h1 {{ color: #1890ff; }}
        h2 {{ color: #262626; background: #fafafa; padding: 8px 12px; border-radius: 4px; }}
        h3 {{ color: #595959; }}
        .context {{ background: #f0f9ff; border-left: 4px solid #1890ff; padding: 12px; margin: 8px 0; }}
        .analysis {{ background: #f6ffed; border-left: 4px solid #52c41a; padding: 12px; margin: 8px 0; }}
        ul li {{ margin: 4px 0; padding: 4px; background: #fff2f0; border-left: 3px solid #ff4d4f; }}
    </style>
</head>
<body>
    <div class="report-header">
        <h1>📚 文献引用合规性检查报告</h1>
    </div>
    <div class="report-meta">
        <p><strong>分析模式：</strong>{task.result['analysis_mode']}</p>
        <p><strong>生成时间：</strong>{task.result['completed_at']}</p>
        <p><strong>发现问题：</strong>{len(task.result['problems']) if 'problems' in task.result else 'N/A'} 个</p>
    </div>
    {task.result['report']}
    <hr style="margin: 30px 0;">
    <p style="text-align: center; color: #999; font-size: 12px;">
        报告由 PaperCheck 文献引用合规性检查工具生成
    </p>
</body>
</html>
"""
    
    # 创建响应
    return Response(
        html_content,
        mimetype='text/html',
        headers={
            'Content-Disposition': f'attachment; filename={filename}'
        }
    )


def export_report_as_txt(task, timestamp):
    """导出为TXT格式"""
    filename = f"citation_report_{task.result['analysis_mode']}_{timestamp}.txt"
    
    # 使用BeautifulSoup提取纯文本
    from bs4 import BeautifulSoup
    
    # 解析HTML报告
    soup = BeautifulSoup(task.result['report'], 'html.parser')
    
    # 提取文本内容
    text_content = soup.get_text()
    
    # 创建完整的TXT内容
    txt_content = f"""📚 文献引用合规性检查报告

分析模式：{task.result['analysis_mode']}
生成时间：{task.result['completed_at']}
发现问题：{len(task.result['problems']) if 'problems' in task.result else 'N/A'} 个

{'='*50}

{text_content}

{'='*50}

报告由 PaperCheck 文献引用合规性检查工具生成
"""
    
    # 创建响应
    return Response(
        txt_content,
        mimetype='text/plain',
        headers={
            'Content-Disposition': f'attachment; filename={filename}'
        }
    )


def export_report_as_pdf(task, timestamp):
    """导出为PDF格式"""
    filename = f"citation_report_{task.result['analysis_mode']}_{timestamp}.pdf"
    
    try:
        # 尝试导入weasyprint
        from weasyprint import HTML, CSS
    except ImportError as e:
        # 如果导入失败，返回友好的错误信息
        error_msg = "PDF导出功能不可用：缺少必要的系统依赖库。请使用HTML或TXT格式导出。"
        print(f"PDF库导入错误: {str(e)}")
        
        # 返回HTML格式的错误提示
        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>PDF导出不可用</title>
</head>
<body>
    <h1>PDF导出功能不可用</h1>
    <p>错误信息: {error_msg}</p>
    <p>建议解决方案:</p>
    <ul>
        <li>请选择HTML或TXT格式进行导出</li>
        <li>如果您需要PDF格式，请联系系统管理员安装所需的依赖库</li>
    </ul>
    <hr>
    <h2>报告内容预览:</h2>
    <div style="max-height: 300px; overflow: auto; border: 1px solid #ccc; padding: 10px;">
        {task.result['report'][:1000]}...
    </div>
    <p><a href="/api/document/{task.task_id}/export?format=html">点击此处下载HTML格式报告</a></p>
    <p><a href="/api/document/{task.task_id}/export?format=txt">点击此处下载TXT格式报告</a></p>
</body>
</html>
"""
        
        return Response(
            html_content,
            mimetype='text/html',
            headers={
                'Content-Disposition': f'attachment; filename={filename.replace(".pdf", "_unavailable.html")}'
            }
        )
    
    try:
        # 创建完整的HTML文件（用于PDF转换）
        html_content = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>文献引用合规性检查报告</title>
    <style>
        @page {{ margin: 2cm; }}
        body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; margin: 2cm; line-height: 1.6; }}
        .report-header {{ text-align: center; margin-bottom: 30px; }}
        .report-meta {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
        h1 {{ color: #1890ff; }}
        h2 {{ color: #262626; background: #fafafa; padding: 8px 12px; border-radius: 4px; }}
        h3 {{ color: #595959; }}
        .context {{ background: #f0f9ff; border-left: 4px solid #1890ff; padding: 12px; margin: 8px 0; }}
        .analysis {{ background: #f6ffed; border-left: 4px solid #52c41a; padding: 12px; margin: 8px 0; }}
        ul li {{ margin: 4px 0; padding: 4px; background: #fff2f0; border-left: 3px solid #ff4d4f; }}
    </style>
</head>
<body>
    <div class="report-header">
        <h1>📚 文献引用合规性检查报告</h1>
    </div>
    <div class="report-meta">
        <p><strong>分析模式：</strong>{task.result['analysis_mode']}</p>
        <p><strong>生成时间：</strong>{task.result['completed_at']}</p>
        <p><strong>发现问题：</strong>{len(task.result['problems']) if 'problems' in task.result else 'N/A'} 个</p>
    </div>
    {task.result['report']}
    <hr style="margin: 30px 0;">
    <p style="text-align: center; color: #999; font-size: 12px;">
        报告由 PaperCheck 文献引用合规性检查工具生成
    </p>
</body>
</html>
"""
        
        # 转换为PDF
        pdf_content = HTML(string=html_content).write_pdf()
        
        # 创建响应
        return Response(
            pdf_content,
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename={filename}'
            }
        )
    except Exception as e:
        # 如果PDF生成失败，返回错误信息
        error_msg = f"PDF导出失败: {str(e)}。系统可能缺少必要的依赖库。"
        print(f"PDF生成错误: {error_msg}")
        
        # 作为备选方案，返回HTML格式并提示用户
        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>导出错误</title>
</head>
<body>
    <h1>PDF导出失败</h1>
    <p>错误信息: {error_msg}</p>
    <p>建议: 请选择HTML或TXT格式进行导出，或者联系系统管理员安装PDF生成所需的依赖。</p>
    <hr>
    <h2>原始报告内容:</h2>
    {task.result['report']}
</body>
</html>
"""
        
        return Response(
            html_content,
            mimetype='text/html',
            headers={
                'Content-Disposition': f'attachment; filename={filename.replace(".pdf", "_error.html")}'
            }
        )

@app.route('/api/document/<task_id>/cancel', methods=['POST'])
def cancel_analysis(task_id):
    """取消分析任务"""
    try:
        if task_id not in analysis_tasks:
            return jsonify({'error': '任务不存在'}), 404
        
        task = analysis_tasks[task_id]
        
        if task.status == 'completed':
            return jsonify({'error': '任务已完成，无法取消'}), 400
        
        task.status = 'cancelled'
        task.error = '用户取消了分析任务'
        
        return jsonify({
            'task_id': task_id,
            'message': '分析任务已取消'
        })
        
    except Exception as e:
        return jsonify({'error': f'取消任务失败: {str(e)}'}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """健康检查接口"""
    return jsonify({
        'status': 'ok',
        'message': '文献分析API服务正常',
        'timestamp': datetime.now().isoformat()
    })

@app.route('/api/images/<filename>')
def serve_image(filename):
    """提供图片文件访问"""
    try:
        image_path = os.path.join('temp_images', filename)
        if os.path.exists(image_path):
            return send_file(image_path)
        else:
            return jsonify({'error': '图片文件不存在'}), 404
    except Exception as e:
        return jsonify({'error': f'获取图片失败: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5001)
