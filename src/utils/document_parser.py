"""
文档解析器 - 将Word文档转换为HTML格式，保留格式特征
"""
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import re
from typing import List, Dict, Any

class DocumentParser:
    """文档解析器类"""
    
    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.doc = docx.Document(doc_path)
        self.paragraphs_data = []
        self._parse_document()
    
    def _parse_document(self):
        """解析文档内容"""
        para_index = 0
        for i, paragraph in enumerate(self.doc.paragraphs):
            para_data = {
                'index': para_index,
                'text': paragraph.text,
                'style': self._get_paragraph_style(paragraph),
                'runs': self._parse_runs(paragraph.runs),
                'type': 'paragraph',
                'doc_index': i  # 在原始文档中的索引
            }
            self.paragraphs_data.append(para_data)
            para_index += 1
    
    def _get_paragraph_style(self, paragraph) -> Dict[str, Any]:
        """获取段落样式"""
        style = {}
        
        # 段落对齐方式
        try:
            if paragraph.alignment:
                alignment_map = {
                    0: 'left',
                    1: 'center',
                    2: 'right',
                    3: 'justify'
                }
                style['text-align'] = alignment_map.get(paragraph.alignment, 'left')
        except (AttributeError, TypeError):
            pass
        
        # 段落间距
        try:
            if paragraph.paragraph_format.space_before:
                style['margin-top'] = f"{paragraph.paragraph_format.space_before.pt}pt"
        except (AttributeError, TypeError):
            pass
            
        try:
            if paragraph.paragraph_format.space_after:
                style['margin-bottom'] = f"{paragraph.paragraph_format.space_after.pt}pt"
        except (AttributeError, TypeError):
            pass
        
        # 首行缩进
        try:
            if paragraph.paragraph_format.first_line_indent:
                style['text-indent'] = f"{paragraph.paragraph_format.first_line_indent.pt}pt"
        except (AttributeError, TypeError):
            pass
        
        return style
    
    def _parse_runs(self, runs) -> List[Dict[str, Any]]:
        """解析文本运行（run）"""
        runs_data = []
        
        for run in runs:
            run_data = {
                'text': run.text,
                'style': self._get_run_style(run)
            }
            runs_data.append(run_data)
        
        return runs_data
    
    def _get_run_style(self, run) -> Dict[str, Any]:
        """获取文本运行样式"""
        style = {}
        
        # 字体
        try:
            if run.font.name:
                style['font-family'] = run.font.name
        except (AttributeError, TypeError):
            pass
        
        # 字体大小
        try:
            if run.font.size:
                style['font-size'] = f"{run.font.size.pt}pt"
        except (AttributeError, TypeError):
            pass
        
        # 加粗
        try:
            if run.bold:
                style['font-weight'] = 'bold'
        except (AttributeError, TypeError):
            pass
        
        # 斜体
        try:
            if run.italic:
                style['font-style'] = 'italic'
        except (AttributeError, TypeError):
            pass
        
        # 下划线
        try:
            if run.underline:
                style['text-decoration'] = 'underline'
        except (AttributeError, TypeError):
            pass
        
        # 字体颜色
        if run.font.color.rgb:
            try:
                rgb = run.font.color.rgb
                # RGBColor对象可能有不同的属性名
                if hasattr(rgb, 'red'):
                    style['color'] = f"rgb({rgb.red}, {rgb.green}, {rgb.blue})"
                elif hasattr(rgb, 'r'):
                    style['color'] = f"rgb({rgb.r}, {rgb.g}, {rgb.b})"
                else:
                    # 如果是整数值，转换为RGB
                    rgb_int = int(rgb) if isinstance(rgb, (int, str)) else rgb
                    r = (rgb_int >> 16) & 0xFF
                    g = (rgb_int >> 8) & 0xFF
                    b = rgb_int & 0xFF
                    style['color'] = f"rgb({r}, {g}, {b})"
            except (AttributeError, TypeError, ValueError):
                # 如果无法解析颜色，跳过
                pass
        
        # 高亮颜色
        try:
            if run.font.highlight_color:
                highlight_colors = {
                    WD_COLOR_INDEX.YELLOW: '#ffff00',
                    WD_COLOR_INDEX.BRIGHT_GREEN: '#00ff00',
                    WD_COLOR_INDEX.TURQUOISE: '#40e0d0',
                    WD_COLOR_INDEX.PINK: '#ffc0cb',
                    WD_COLOR_INDEX.BLUE: '#0000ff',
                    WD_COLOR_INDEX.RED: '#ff0000',
                    WD_COLOR_INDEX.DARK_BLUE: '#00008b',
                    WD_COLOR_INDEX.TEAL: '#008080',
                    WD_COLOR_INDEX.GREEN: '#008000',
                    WD_COLOR_INDEX.VIOLET: '#8b00ff',
                    WD_COLOR_INDEX.DARK_RED: '#8b0000',
                    WD_COLOR_INDEX.DARK_YELLOW: '#8b8b00',
                    WD_COLOR_INDEX.GRAY_25: '#c0c0c0',
                    WD_COLOR_INDEX.GRAY_50: '#808080'
                }
                if run.font.highlight_color in highlight_colors:
                    style['background-color'] = highlight_colors[run.font.highlight_color]
        except (AttributeError, TypeError):
            # 如果无法处理高亮颜色，跳过
            pass
        
        return style
    
    def to_html(self) -> str:
        """转换为HTML格式，保持原始顺序"""
        html_parts = ['<div class="document-content">']
        
        # 创建一个包含所有内容（段落、表格和图片）的列表，并按原始顺序排列
        content_items = []
        
        # 收集段落
        for i, paragraph in enumerate(self.doc.paragraphs):
            content_items.append(('paragraph', i, paragraph, i))
            
        # 收集表格
        for i, table in enumerate(self.doc.tables):
            # 获取表格在文档中的位置
            table_position = table._element.getparent().index(table._element)
            content_items.append(('table', i, table, table_position))
            
        # 收集图片
        for i, shape in enumerate(self.doc.inline_shapes):
            # 获取图片在文档中的位置
            # 简化处理，将所有图片放在文档末尾
            content_items.append(('image', i, shape, len(self.doc.paragraphs) + i))
            
        # 按照在文档中的实际顺序排序
        content_items.sort(key=lambda x: x[3])
        
        # 按顺序处理所有元素
        for element_type, original_index, element, doc_index in content_items:
            if element_type == 'paragraph':
                paragraph = element
                # 查找对应的段落数据
                para_data = None
                for pd in self.paragraphs_data:
                    if pd.get('doc_index') == original_index:
                        para_data = pd
                        break
                
                if para_data:
                    # 构建段落HTML
                    para_style = self._style_dict_to_css(para_data['style'])
                    html_parts.append(f'<p class="paragraph" data-para-index="{para_data["index"]}" style="{para_style}">')
                    
                    # 处理每个run
                    for run_data in para_data['runs']:
                        if run_data['text']:
                            run_style = self._style_dict_to_css(run_data['style'])
                            escaped_text = self._escape_html(run_data['text'])
                            html_parts.append(f'<span style="{run_style}">{escaped_text}</span>')
                    
                    html_parts.append('</p>')
            elif element_type == 'table':
                table = element
                html_parts.append(self._table_to_html(table))
            else:  # image
                shape = element
                # 处理图片
                image_html = self._image_to_html(shape)
                if image_html:
                    html_parts.append(image_html)
        
        html_parts.append('</div>')
        
        return ''.join(html_parts)
    
    def _style_dict_to_css(self, style_dict: Dict[str, Any]) -> str:
        """将样式字典转换为CSS字符串"""
        if not style_dict:
            return ''
        
        css_parts = []
        for key, value in style_dict.items():
            css_parts.append(f"{key}: {value}")
        
        return '; '.join(css_parts)
    
    def _escape_html(self, text: str) -> str:
        """转义HTML特殊字符"""
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&#x27;'))
    
    def get_paragraphs_with_positions(self) -> List[Dict[str, Any]]:
        """获取带位置信息的段落列表，保持原始顺序"""
        # 重新解析文档以获取正确的顺序
        paragraphs_data = []
        
        # 创建一个包含所有内容（段落、表格和图片）的列表，并按原始顺序排列
        content_items = []
        
        # 收集段落
        for i, paragraph in enumerate(self.doc.paragraphs):
            content_items.append(('paragraph', i, paragraph, i))
            
        # 收集表格
        for i, table in enumerate(self.doc.tables):
            # 获取表格在文档中的位置
            table_position = table._element.getparent().index(table._element)
            content_items.append(('table', i, table, table_position))
            
        # 收集图片
        for i, shape in enumerate(self.doc.inline_shapes):
            # 获取图片在文档中的位置
            # 简化处理，将所有图片放在文档末尾
            content_items.append(('image', i, shape, len(self.doc.paragraphs) + i))
            
        # 按照在文档中的实际顺序排序
        content_items.sort(key=lambda x: x[3])
        
        # 按顺序处理所有元素
        para_index = 0
        for element_type, original_index, element, doc_index in content_items:
            if element_type == 'paragraph':
                paragraph = element
                para_info = {
                    'index': para_index,
                    'text': paragraph.text,
                    'length': len(paragraph.text),
                    'style': self._get_paragraph_style(paragraph)
                }
                paragraphs_data.append(para_info)
                para_index += 1
            elif element_type == 'table':
                table = element
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():  # 只处理非空段落
                                para_info = {
                                    'index': para_index,
                                    'text': paragraph.text,
                                    'length': len(paragraph.text),
                                    'style': self._get_paragraph_style(paragraph),
                                    'is_table_cell': True
                                }
                                paragraphs_data.append(para_info)
                                para_index += 1
            else:  # image
                shape = element
                # 对于图片，我们添加一个占位符段落
                para_info = {
                    'index': para_index,
                    'text': '[图片]',
                    'length': 4,
                    'style': {},
                    'is_image': True
                }
                paragraphs_data.append(para_info)
                para_index += 1
        
        return paragraphs_data
    
    def find_text_positions(self, search_text: str) -> List[Dict[str, Any]]:
        """查找文本在文档中的位置"""
        positions = []
        
        for para_data in self.paragraphs_data:
            text = para_data['text']
            start = 0
            
            while True:
                pos = text.find(search_text, start)
                if pos == -1:
                    break
                
                positions.append({
                    'paragraph_index': para_data['index'],
                    'start_pos': pos,
                    'end_pos': pos + len(search_text),
                    'context': text
                })
                
                start = pos + 1
        
        return positions
    
    def _get_table_style(self, table) -> Dict[str, Any]:
        """获取表格样式"""
        style = {}
        
        try:
            # 获取表格的底层XML元素
            tbl = table._tbl
            tbl_pr = tbl.tblPr
            
            # 获取背景色信息
            tbl_shd = tbl_pr.xpath('.//w:shd')
            if tbl_shd:
                shd = tbl_shd[0]
                fill_color = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if fill_color:
                    style['background-color'] = f"#{fill_color}"
                    
            # 获取边框信息
            tbl_borders = tbl_pr.xpath('.//w:tblBorders')
            if tbl_borders:
                # 简化处理，只应用一个边框样式
                style['border'] = "1px solid #000"
        except (AttributeError, TypeError):
            pass
        
        return style
    
    def _get_table_cell_style(self, cell) -> Dict[str, Any]:
        """获取表格单元格样式"""
        style = {}
        
        # 获取单元格的宽度
        try:
            if cell.width:
                # 转换EMU单位到像素（1英寸=914400 EMU, 1英寸≈96像素）
                width_px = int(cell.width / 914400 * 96)
                style['width'] = f"{width_px}px"
        except (AttributeError, TypeError):
            pass
            
        # 获取单元格背景色
        try:
            # 获取单元格的底层XML元素
            tc = cell._tc
            tc_pr = tc.tcPr
            
            # 获取背景色信息
            tc_shd = tc_pr.xpath('.//w:shd')
            if tc_shd:
                shd = tc_shd[0]
                fill_color = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if fill_color:
                    style['background-color'] = f"#{fill_color}"
        except (AttributeError, TypeError):
            pass
        
        return style
    
    def _get_table_row_style(self, row) -> Dict[str, Any]:
        """获取表格行样式"""
        style = {}
        
        # 获取行高度
        try:
            if row.height:
                # 转换EMU单位到像素
                height_px = int(row.height / 914400 * 96)
                style['height'] = f"{height_px}px"
        except (AttributeError, TypeError):
            pass
        
        return style
    
    def _table_to_html(self, table) -> str:
        """将表格转换为HTML格式"""
        # 获取表格样式
        table_style = self._get_table_style(table)
        base_table_style = "border-collapse: collapse; width: 100%; margin: 10px 0;"
        table_style_str = self._style_dict_to_css(table_style)
        full_table_style = base_table_style + (f"; {table_style_str}" if table_style_str else "")
        
        html_parts = [f'<table style="{full_table_style}">']
        
        for i, row in enumerate(table.rows):
            # 获取行样式
            row_style = self._get_table_row_style(row)
            row_style_str = self._style_dict_to_css(row_style)
            html_parts.append(f'<tr style="{row_style_str}">')
            
            for j, cell in enumerate(row.cells):
                # 获取单元格样式
                cell_style = self._get_table_cell_style(cell)
                
                # 获取单元格内容（包括格式）
                cell_content_parts = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        # 查找对应的段落数据
                        para_data = None
                        for p in self.paragraphs_data:
                            if p['text'] == paragraph.text and p.get('is_table_cell'):
                                para_data = p
                                break
                        
                        if para_data:
                            # 处理段落中的每个run
                            for run in paragraph.runs:
                                if run.text:
                                    run_style = self._get_run_style(run)
                                    run_style_str = self._style_dict_to_css(run_style)
                                    escaped_text = self._escape_html(run.text)
                                    cell_content_parts.append(f'<span style="{run_style_str}">{escaped_text}</span>')
                        else:
                            # 备用方案：直接处理文本
                            escaped_text = self._escape_html(paragraph.text)
                            cell_content_parts.append(escaped_text)
                
                # 合并单元格内容
                cell_content = ''.join(cell_content_parts)
                
                # 单元格基本样式
                base_cell_style = "border: 1px solid #ddd; padding: 8px; vertical-align: top;"
                cell_style_str = self._style_dict_to_css(cell_style)
                full_cell_style = base_cell_style + (f"; {cell_style_str}" if cell_style_str else "")
                
                html_parts.append(f'<td style="{full_cell_style}">{cell_content}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</table>')
        return ''.join(html_parts)
    
    def _get_table_cell_style(self, cell) -> Dict[str, Any]:
        """获取表格单元格样式"""
        style = {}
        
        # 获取单元格的宽度
        try:
            if cell.width:
                # 转换EMU单位到像素（1英寸=914400 EMU, 1英寸≈96像素）
                width_px = int(cell.width / 914400 * 96)
                style['width'] = f"{width_px}px"
        except (AttributeError, TypeError):
            pass
            
        # 获取单元格背景色
        try:
            # 获取单元格的底层XML元素
            tc = cell._tc
            tc_pr = tc.tcPr
            
            # 获取背景色信息
            tc_shd = tc_pr.xpath('.//w:shd')
            if tc_shd:
                shd = tc_shd[0]
                fill_color = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if fill_color:
                    style['background-color'] = f"#{fill_color}"
        except (AttributeError, TypeError):
            pass
        
        return style
    
    def _get_table_row_style(self, row) -> Dict[str, Any]:
        """获取表格行样式"""
        style = {}
        
        # 获取行高度
        try:
            if row.height:
                # 转换EMU单位到像素
                height_px = int(row.height / 914400 * 96)
                style['height'] = f"{height_px}px"
        except (AttributeError, TypeError):
            pass
        
        return style
    
    def _table_to_html(self, table) -> str:
        """将表格转换为HTML格式"""
        # 获取表格样式
        table_style = self._get_table_style(table)
        base_table_style = "border-collapse: collapse; width: 100%; margin: 10px 0;"
        table_style_str = self._style_dict_to_css(table_style)
        full_table_style = base_table_style + (f"; {table_style_str}" if table_style_str else "")
        
        html_parts = [f'<table style="{full_table_style}">']
        
        for i, row in enumerate(table.rows):
            # 获取行样式
            row_style = self._get_table_row_style(row)
            row_style_str = self._style_dict_to_css(row_style)
            html_parts.append(f'<tr style="{row_style_str}">')
            
            for j, cell in enumerate(row.cells):
                # 获取单元格样式
                cell_style = self._get_table_cell_style(cell)
                
                # 获取单元格内容（包括格式）
                cell_content_parts = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        # 处理段落中的每个run
                        for run in paragraph.runs:
                            if run.text:
                                run_style = self._get_run_style(run)
                                run_style_str = self._style_dict_to_css(run_style)
                                escaped_text = self._escape_html(run.text)
                                cell_content_parts.append(f'<span style="{run_style_str}">{escaped_text}</span>')
                
                # 合并单元格内容
                cell_content = ''.join(cell_content_parts)
                
                # 单元格基本样式
                base_cell_style = "border: 1px solid #ddd; padding: 8px; vertical-align: top;"
                cell_style_str = self._style_dict_to_css(cell_style)
                full_cell_style = base_cell_style + (f"; {cell_style_str}" if cell_style_str else "")
                
                html_parts.append(f'<td style="{full_cell_style}">{cell_content}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</table>')
        return ''.join(html_parts)
    
    def _image_to_html(self, shape) -> str:
        """将图片转换为HTML格式"""
        try:
            # 检查是否为图片
            if not hasattr(shape, 'graphic') or not hasattr(shape.graphic, 'graphicData'):
                return ''
                
            # 获取图片
            image_part = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_data = self.doc.part.related_parts[image_part]._blob
            
            # 生成唯一的图片文件名
            import hashlib
            image_hash = hashlib.md5(image_data).hexdigest()[:8]
            image_filename = f"image_{image_hash}.png"
            
            # 保存图片到临时目录
            import os
            temp_dir = "temp_images"
            if not os.path.exists(temp_dir):
                os.makedirs(temp_dir)
                
            image_path = os.path.join(temp_dir, image_filename)
            with open(image_path, "wb") as f:
                f.write(image_data)
            
            # 获取图片尺寸信息
            try:
                width = shape.width // 914400 * 96  # 转换EMUs到像素
                height = shape.height // 914400 * 96
            except:
                width = 400
                height = 300
            
            # 生成HTML，使用API端点提供图片访问
            return f'<div style="text-align: center; margin: 10px 0;"><img src="/api/images/{image_filename}" alt="文档图片" style="max-width: 100%; height: auto;" width="{width}" height="{height}"></div>'
        except Exception as e:
            # 如果无法处理图片，返回一个占位符
            return '<div style="text-align: center; margin: 10px 0; padding: 20px; border: 1px dashed #ccc; color: #666;">[图片]</div>'
    
    def get_paragraph_by_index(self, index: int) -> Dict[str, Any]:
        """根据索引获取段落"""
        # 重新解析文档以获取正确的顺序
        paragraphs_data = self.get_paragraphs_with_positions()
        if 0 <= index < len(paragraphs_data):
            return paragraphs_data[index]
        return None
