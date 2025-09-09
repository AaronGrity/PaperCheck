"""
位置跟踪器 - 精确定位引用和参考文献在文档中的位置
"""
import docx
import re
from typing import List, Dict, Any, Tuple

class PositionTracker:
    """位置跟踪器类"""
    
    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.doc = docx.Document(doc_path)
        self.paragraphs = []
        self.references_start_index = None
        self._parse_document()
    
    def _parse_document(self):
        """解析文档，建立位置索引，保持原始顺序"""
        # 创建一个包含所有内容（段落、表格和图片）的列表，并按原始顺序排列
        content_items = []
        
        # 收集段落
        for i, paragraph in enumerate(self.doc.paragraphs):
            content_items.append(('paragraph', i, paragraph, i))
            
        # 收集表格
        for i, table in enumerate(self.doc.tables):
            # 获取表格在文档中的位置
            table_position = table._element.getparent().index(table._element)
            content_items.append(('table', i, table, table_position))
            
        # 收集图片
        for i, shape in enumerate(self.doc.inline_shapes):
            # 获取图片在文档中的位置
            # 简化处理，将所有图片放在文档末尾
            content_items.append(('image', i, shape, len(self.doc.paragraphs) + i))
            
        # 按照在文档中的实际顺序排序
        content_items.sort(key=lambda x: x[3])
        
        # 按顺序处理所有元素
        para_index = 0
        for element_type, original_index, element, doc_index in content_items:
            if element_type == 'paragraph':
                paragraph = element
                para_info = {
                    'index': para_index,
                    'text': paragraph.text,
                    'length': len(paragraph.text),
                    'type': 'paragraph'
                }
                self.paragraphs.append(para_info)
                para_index += 1
                
                # 查找参考文献开始位置
                if self.references_start_index is None:
                    if '参考文献' in paragraph.text or 'References' in paragraph.text:
                        self.references_start_index = para_index - 1
            elif element_type == 'table':
                table = element
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():  # 只处理非空段落
                                para_info = {
                                    'index': para_index,
                                    'text': paragraph.text,
                                    'length': len(paragraph.text),
                                    'type': 'table_cell'
                                }
                                self.paragraphs.append(para_info)
                                para_index += 1
            else:  # image
                # 对于图片，我们添加一个占位符段落
                para_info = {
                    'index': para_index,
                    'text': '[图片]',
                    'length': 4,
                    'type': 'image'
                }
                self.paragraphs.append(para_info)
                para_index += 1
    
    def find_citation_positions(self, citation: str) -> List[Dict[str, Any]]:
        """查找引用在文档中的所有位置"""
        positions = []
        
        # 优先在正文中查找（参考文献之前）
        search_range = self.paragraphs[:self.references_start_index] if self.references_start_index else self.paragraphs
        
        for para in search_range:
            text = para['text']
            # 查找所有匹配的位置
            for match in re.finditer(re.escape(citation), text):
                position = {
                    'paragraph_index': para['index'],
                    'start_pos': match.start(),
                    'end_pos': match.end(),
                    'context': self._get_context(para['index'], match.start(), match.end()),
                    'in_references': False
                }
                positions.append(position)
        
        # 如果在正文中没找到，再在参考文献中查找
        if not positions and self.references_start_index:
            ref_range = self.paragraphs[self.references_start_index:]
            for para in ref_range:
                text = para['text']
                for match in re.finditer(re.escape(citation), text):
                    position = {
                        'paragraph_index': para['index'],
                        'start_pos': match.start(),
                        'end_pos': match.end(),
                        'context': self._get_context(para['index'], match.start(), match.end()),
                        'in_references': True
                    }
                    positions.append(position)
        
        return positions
    
    def find_reference_positions(self, reference_text: str) -> List[Dict[str, Any]]:
        """查找参考文献在文档中的位置"""
        positions = []
        
        # 在参考文献部分查找
        if self.references_start_index:
            ref_range = self.paragraphs[self.references_start_index:]
            
            for para in ref_range:
                text = para['text']
                # 检查是否是参考文献条目
                if text.strip() and (text.strip().startswith('[') or reference_text in text):
                    position = {
                        'paragraph_index': para['index'],
                        'start_pos': 0,
                        'end_pos': len(text),
                        'context': text,
                        'in_references': True
                    }
                    positions.append(position)
        
        return positions
    
    def _get_context(self, para_index: int, start_pos: int, end_pos: int, context_length: int = 100) -> Dict[str, Any]:
        """获取指定位置的上下文信息"""
        if para_index >= len(self.paragraphs):
            return {}
        
        para = self.paragraphs[para_index]
        text = para['text']
        
        # 获取当前段落的上下文
        context_start = max(0, start_pos - context_length)
        context_end = min(len(text), end_pos + context_length)
        context_text = text[context_start:context_end]
        
        # 获取前后段落作为扩展上下文
        extended_context = []
        
        # 前一段落
        if para_index > 0 and self.paragraphs[para_index - 1]['text'].strip():
            extended_context.append(self.paragraphs[para_index - 1]['text'].strip())
        
        # 当前段落
        extended_context.append(text)
        
        # 后一段落
        if para_index < len(self.paragraphs) - 1 and self.paragraphs[para_index + 1]['text'].strip():
            extended_context.append(self.paragraphs[para_index + 1]['text'].strip())
        
        return {
            'paragraph_text': text,
            'context_text': context_text,
            'extended_context': ' '.join(extended_context),
            'relative_start': start_pos - context_start,
            'relative_end': end_pos - context_start
        }
    
    def find_all_citations(self) -> List[Dict[str, Any]]:
        """查找文档中的所有引用"""
        citation_pattern = r'\[\d+(?:-\d+)?\]'
        all_citations = []
        
        # 只在正文中查找（参考文献之前）
        search_range = self.paragraphs[:self.references_start_index] if self.references_start_index else self.paragraphs
        
        for para in search_range:
            text = para['text']
            for match in re.finditer(citation_pattern, text):
                citation = match.group()
                
                # 展开范围引用
                expanded_citations = self._expand_citation_range(citation)
                
                for expanded_citation in expanded_citations:
                    citation_info = {
                        'citation': expanded_citation,
                        'original_citation': citation,
                        'paragraph_index': para['index'],
                        'start_pos': match.start(),
                        'end_pos': match.end(),
                        'context': self._get_context(para['index'], match.start(), match.end())
                    }
                    all_citations.append(citation_info)
        
        return all_citations
    
    def _expand_citation_range(self, citation: str) -> List[str]:
        """展开范围引用为单个引用列表"""
        if '-' not in citation:
            return [citation]
        
        match = re.match(r'\[(\d+)-(\d+)\]', citation)
        if not match:
            return [citation]
        
        start = int(match.group(1))
        end = int(match.group(2))
        
        return [f'[{i}]' for i in range(start, end + 1)]
    
    def get_paragraph_by_position(self, para_index: int) -> Dict[str, Any]:
        """根据段落索引获取段落信息"""
        if 0 <= para_index < len(self.paragraphs):
            return self.paragraphs[para_index]
        return None
    
    def find_text_in_range(self, search_text: str, start_para: int = 0, end_para: int = None) -> List[Dict[str, Any]]:
        """在指定段落范围内查找文本"""
        if end_para is None:
            end_para = len(self.paragraphs)
        
        positions = []
        search_range = self.paragraphs[start_para:end_para]
        
        for i, para in enumerate(search_range):
            actual_index = start_para + i
            text = para['text']
            
            for match in re.finditer(re.escape(search_text), text):
                position = {
                    'paragraph_index': actual_index,
                    'start_pos': match.start(),
                    'end_pos': match.end(),
                    'context': self._get_context(actual_index, match.start(), match.end())
                }
                positions.append(position)
        
        return positions
    
    def get_document_structure(self) -> Dict[str, Any]:
        """获取文档结构信息"""
        return {
            'total_paragraphs': len(self.paragraphs),
            'references_start_index': self.references_start_index,
            'body_paragraphs': self.references_start_index if self.references_start_index else len(self.paragraphs),
            'reference_paragraphs': len(self.paragraphs) - self.references_start_index if self.references_start_index else 0
        }
    
    def validate_citation_positions(self, citations: List[str]) -> Dict[str, List[Dict[str, Any]]]:
        """验证引用位置信息"""
        validation_result = {}
        
        for citation in citations:
            positions = self.find_citation_positions(citation)
            validation_result[citation] = positions
        
        return validation_result
