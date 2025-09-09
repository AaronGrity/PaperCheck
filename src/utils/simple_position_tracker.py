"""
简化的位置跟踪器 - 专注引用定位
"""
import docx
import re
from typing import List, Dict

class SimplePositionTracker:
    """简化的位置跟踪器"""
    
    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.doc = docx.Document(doc_path)
        self.paragraphs = [p.text for p in self.doc.paragraphs]
    
    def find_citation_positions(self, citation: str) -> List[Dict]:
        """查找引用在文档中的位置"""
        positions = []
        
        for para_idx, para_text in enumerate(self.paragraphs):
            # 查找所有匹配的位置
            for match in re.finditer(re.escape(citation), para_text):
                positions.append({
                    'paragraph_index': para_idx,
                    'start_pos': match.start(),
                    'end_pos': match.end(),
                    'context': para_text,
                    'preview': self._get_preview(para_text, match.start(), match.end())
                })
        
        return positions
    
    def _get_preview(self, text: str, start: int, end: int, context_length: int = 50) -> str:
        """获取引用周围的预览文本"""
        preview_start = max(0, start - context_length)
        preview_end = min(len(text), end + context_length)
        
        preview = text[preview_start:preview_end]
        
        # 添加省略号
        if preview_start > 0:
            preview = "..." + preview
        if preview_end < len(text):
            preview = preview + "..."
            
        return preview
    
    def get_all_citations(self) -> List[str]:
        """获取文档中的所有引用"""
        citations = set()
        citation_pattern = r'\[\d+(?:-\d+)?\]'
        
        for para_text in self.paragraphs:
            matches = re.findall(citation_pattern, para_text)
            for match in matches:
                # 展开范围引用
                if '-' in match:
                    range_match = re.match(r'\[(\d+)-(\d+)\]', match)
                    if range_match:
                        start_num = int(range_match.group(1))
                        end_num = int(range_match.group(2))
                        for i in range(start_num, end_num + 1):
                            citations.add(f'[{i}]')
                else:
                    citations.add(match)
        
        return sorted(list(citations), key=lambda x: int(re.search(r'\[(\d+)\]', x).group(1)))
    
    def get_document_info(self) -> Dict:
        """获取文档基本信息"""
        return {
            'total_paragraphs': len(self.paragraphs),
            'total_characters': sum(len(p) for p in self.paragraphs),
            'citations': self.get_all_citations()
        }
